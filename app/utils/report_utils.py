import matplotlib
matplotlib.use('Agg')

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
import docx
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from collections import defaultdict
import os
import matplotlib.pyplot as plt
from io import BytesIO
import io
import plotly.graph_objects as go
import base64
from flask import current_app

def gerar_relatorio_completo(top_5_solicitacoes, data_records, total_abertos, total_concluidos, chamados_por_data, pipefy_image, grafana_image_paths, autorResponsavel, numberVersion, cards_data, all_cards_data, start_date, end_date, custom_link1, custom_link2, custom_link3):
    """
    Gera o relatório completo diretamente pelo código, com base nas informações do Pipefy, Grafana e Prometheus.
    """
    doc = Document()

    mes_ano_atual = start_date.strftime("%B")
    adicionar_cabecalho_rodape(doc, mes_ano_atual)

    gerar_pagina_inicial(doc, start_date, end_date)
    doc.add_page_break()

    doc.add_heading('Relatório Gerencial Mensal', 1)
    adicionar_sumario(doc)
    doc.add_page_break()

    doc.add_heading('Objetivo', level=2)
    doc.add_paragraph(
        'Mensalmente, será apresentado o relatório “Run-Book” contendo as informações baseadas na operação dos últimos 30 dias, com as seguintes informações, porém não limitado a:')

    doc.add_paragraph(
        '''
        • Relatório de incidentes ocorrido no período anterior contendo a classificação dos incidentes e SLA atingido
        • Descritivo e tratativa dos incidentes críticos ocorridos no período anterior. 
        • Informativo a respeito dos itens de segurança da informação como a atualização de patchs de segurança e hotfixes. 
        • Resumo dos recursos computacionais utilizados contendo % de consumo – capacidade do ambiente computacional. 
        • Resumo da saúde dos elementos tecnológicos fornecidos nesta solução. 
        • Resumo dos alertas gerados pelo sistema de monitoramento para todos os elementos tecnológicos fornecidos pela ITFácil. 
        • Indicativo de risco e recomendação de tomada de ação 
        • Relatório sobre inventário do parque e acompanhamento da vida útil.'''
    )

    doc.add_paragraph(
        'Adicionalmente ao Run-Book mensal, a ITfácil poderá fornecer a pedido do cliente'
    )

    doc.add_paragraph(
        '''
        • Relatório check-list de execução de rotina de backup diário 
        • Relatório check-list de execução de replicação do site Disaster Recovery diário 
        • Check-list operacional a respeito do funcionamento da infraestrutura e sistemas diário 
        • Relatórios de auditoria (conforme demanda)
        '''
    )

    doc.add_heading('SLA - Solução', level=2)
    doc.add_picture('./app/sla_padrao.jpg', width=Inches(4.0), height=Inches(4.0))


    doc.add_heading('Histórico de versões', level=2)
    gerar_historico_versoes(doc, autorResponsavel, numberVersion)

    doc.add_page_break()
    doc.add_heading('Fatos relevantes', level=2)
    gerar_relatorio_chamados(doc, data_records, start_date, end_date)
    doc.add_page_break()

    doc.add_heading('Quantidade de GMUDs Executadas no Mês', level=2)
    doc.add_paragraph('Neste mês, o time de TI realizou um total de [NÚMERO] GMUDs (Gestão de Mudanças em TI), alinhadas com as necessidades estratégicas e operacionais da empresa. Cada GMUD foi cuidadosamente planejada e implementada com o objetivo de otimizar a infraestrutura tecnológica, melhorar a performance dos sistemas e garantir a continuidade dos serviços. O monitoramento pós-implementação confirmou a eficácia das mudanças.')

    doc.add_heading('Principais Tipos de Solicitações', level=2)
    adicionar_texto_solicitacoes(doc, start_date, end_date, top_5_solicitacoes)
    gerar_grafico_colunas_solicitacoes(doc, top_5_solicitacoes)
    doc.add_page_break()

    doc.add_heading('Gestão de Serviços', level=2)
    doc.add_paragraph('Foi implantado a Gestão de projetos na Ecom, utilizando a ferramenta do Pipefy, com a criação de vários controles que podem ser gerenciados pela equipe Ecom e Itfácil em conjunto, essa ferramenta trouxe mais transparências nas atividades, registro dos passos e evidências além de criar uma base de conhecimento que pode ser consultada a qualquer momento.')
    if pipefy_image:
        doc.add_picture(os.path.join(current_app.config['UPLOAD_FOLDER'], pipefy_image), width=Inches(4.0))
    p = doc.add_paragraph('Acesse o quadro Pipefy aqui: ')
    add_hyperlink(p, 'https://app.pipefy.com/pipes/304582953', 'Quadro - Ecom - Tarefas ITFacil Equipe - Pipefy')

    gerar_classificacao_incidentes(doc, data_records, start_date, end_date)

    doc.add_page_break()
    doc.add_heading('Resumo dos recursos computacionais utilizados contendo % de consumo – capacidade do ambiente computacional.', level=2)
    doc.add_paragraph('Considerando a topologia no ambiente Ecom / Cirion –abaixo, esta em desenvolvimento o  monitoramento com a capacidade de listar os % de consumo dos Host por períodos, promovendo as seguintes informações:')
    if grafana_image_paths:
        for grafana_image_path in grafana_image_paths:
            doc.add_picture(grafana_image_path, width=Inches(6.0))
    #gerar_relatorio_com_dados_prometheus(doc, prometheus_data)
    #gerar_relatorio_com_dados_grafana(doc, grafana_data)

    doc.add_heading('Check-list operacional a respeito do funcionamento da infraestrutura e sistemas diário', level=2)
    p = doc.add_paragraph('Disponível no diretório: ')
    add_hyperlink(p, custom_link1, 'Lista de Check List - Diario')

    doc.add_paragraph('Esta disponível o manual do check-list atualizado, com o processo passo-a-passo e um arquivo zip com todos os check list realizados inclusive com as observações e ocorrências.')

    doc.add_heading('Backup Cirion, acompanhamento diário', level=2)
    doc.add_paragraph('O backup é realizado junto ao checklist diário.')

    doc.add_heading('Controle de aplicativos instalados', level=2)
    p = doc.add_paragraph('Disponível no diretório: ')
    add_hyperlink(p, custom_link2, 'Controle de Aplicativos e Sistemas Instalados')

    doc.add_heading('O Sistema GLPI e o Kaspersky fazem a gestão dos ativos controalndo a parte física e lógica de todos os equipamentos', level=2)
    doc.add_paragraph('A lista completa, bem como o inventário por equipamentos poderá ser acompanhado e gerenciado diretamente no GLPI ou Kaspersky')
    p = doc.add_paragraph('Relatório completo está disponível em: ')
    add_hyperlink(p, custom_link3, 'Inventario de Maquinas - GLPI')

    doc.add_heading('Chamados Referente a TOTVS', level=2)
    gerar_chamados_totvs(doc, data_records, start_date, end_date)

    output_file = os.path.join(os.getcwd(), 'relatorio_runbook.docx')
    doc.save(output_file)

    return output_file

def adicionar_cabecalho_rodape(doc, mes_ano):
    """
    Adiciona cabeçalho com imagens e rodapé com texto nas páginas.
    """
    sections = doc.sections

    for section in sections:
        section.different_first_page_header_footer = True

        header = section.header
        header_paragraph = header.paragraphs[0]

        img_itfacil_path = './app/img_itfacil.png'
        header_paragraph.add_run().add_picture(img_itfacil_path, width=Inches(1.25))

        img_ecom_path = './app/img_ecom.png'
        header_paragraph.add_run().add_picture(img_ecom_path, width=Inches(1.25))

        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rodape = section.footer
        rodape_paragraph = rodape.paragraphs[0]

        rodape_text = f'RUN BOOK - {mes_ano}'

        rodape_run = rodape_paragraph.add_run(rodape_text)
        rodape_run.font.size = Pt(12)

        rodape_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    first_section = sections[0]
    first_section.header.is_linked_to_previous = False
    first_section.footer.is_linked_to_previous = False

def adicionar_sumario(doc):
    """
    Gera um sumário com alinhamento adequado para o documento.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Sumário")
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sumario = [
        ("1. Objetivo", 3),
        ("2. SLA - Solução", 4),
        ("3. Histórico de Versões", 4),
        ("4. Fatos relevantes", 5),
        ("5. Quantidade de GMUDs Executadas no Mês", 6),
        ("6. Principais Tipos de Solicitações", 6),
        ("7. Gestão de Serviços", 7),
        ("8. Resumo dos recursos computacionais utilizados", 8),
        ("9. Check-list operacional", 11),
        ("10. Backup Cirion", 11),
        ("11. Controle de aplicativos instalados", 11),
        ("12. O Sistema GLPI e o Kaspersky", 11),
        ("13. Chamados Referente a TOTVS", 12)
    ]

    for titulo, pagina in sumario:
        paragrafo = doc.add_paragraph()
        paragrafo.add_run(titulo)

        tab_stops = paragrafo.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Pt(450), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        paragrafo.add_run("\t" + str(pagina))

def adicionar_imagens_grafana(doc, images):
    """
    Adiciona até 5 imagens do Grafana no documento.
    """
    for i, image_path in enumerate(images[:5]):
        if image_path:
            doc.add_paragraph(f'Imagem Grafana {i+1}')
            doc.add_picture(image_path, width=Inches(5))
            doc.add_paragraph("")

def gerar_grafico_colunas_solicitacoes(doc, top_5_solicitacoes):
    """
    Gera um gráfico de colunas com as 5 principais solicitações abertas no período selecionado.
    """
    if not top_5_solicitacoes:
        doc.add_paragraph("Nenhuma solicitação encontrada no período selecionado.")
        return

    solicitacoes, quantidades = zip(*top_5_solicitacoes)

    plt.figure(figsize=(10, 4))
    plt.bar(solicitacoes, quantidades, color='skyblue')
    plt.title('Tipos de Solicitações Abertas no Período')
    plt.xlabel('Tipo de Solicitação')
    plt.ylabel('Quantidade')
    plt.xticks(rotation=45, ha='right', fontsize=6)
    plt.tight_layout()
    plt.subplots_adjust(bottom=0.45)

    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)

    doc.add_picture(img, width=Inches(6))

    plt.close()

def adicionar_texto_solicitacoes(doc, start_date, end_date, top_5_solicitacoes):
    """
        Adiciona um texto explicativo sobre as cinco principais solicitações abertas no período.
        """
    if not top_5_solicitacoes:
        doc.add_paragraph("Nenhuma solicitação foi registrada no período selecionado.")
        return

    doc.add_paragraph("As cinco principais solicitações abertas no período de " +
                      f"{start_date.strftime('%d/%m/%Y')} a {end_date.strftime('%d/%m/%Y')} foram:")

    for solicitacao, quantidade in top_5_solicitacoes:
        doc.add_paragraph(f"- {solicitacao}: {quantidade} solicitações abertas.")

    doc.add_paragraph("Essas solicitações representam as principais demandas enfrentadas pelo setor durante o período.")

    doc.add_paragraph("O gráfico abaixo apresenta uma visualização das principais solicitações e sua distribuição.")

def gerar_classificacao_incidentes(doc, cards, start_date, end_date):
    """
    Gera a seção 'Classificação dos Incidentes' com gráfico de pizza.
    """
    tipos_solicitacao = defaultdict(int)

    start_date = start_date.date() if isinstance(start_date, datetime) else start_date
    end_date = end_date.date() if isinstance(end_date, datetime) else end_date

    for card in cards:
        card_created_date = card['Created at'].date()
        card_finished_date = card['Finished at'].date()
        current_phase = card.get('Current phase')
        tipo_solicitacao = card.get('Escolha o tipo de solicitação')
        if start_date <= card_created_date <= end_date:
            if tipo_solicitacao:
                tipos_solicitacao[tipo_solicitacao] += 1
            card_title = card.get('Title', 'Descreva sua Solicitação')
            if 'TOTVS' in card_title.upper() or 'DATASUL' in card_title.upper():
                tipos_solicitacao['TOTVS'] += 1

    if tipos_solicitacao:
        gerar_grafico_pizza(doc, tipos_solicitacao)
    else:
        doc.add_paragraph("Nenhum incidente encontrado para classificar.")

    doc.add_paragraph('Classificação dos Incidentes:')
    doc.add_paragraph(f'A classificação dos chamados do setor de infraestrutura para o periodo de {start_date.strftime("%d/%m/%Y")} a {end_date.strftime("%d/%m/%Y")} foi dividida da seguinte forma: Os incidentes, normalmente associados a falhas de sistema ou indisponibilidade de recursos, foram tratados com prioridade para minimizar o impacto nas operações. As solicitações de serviço envolveram demandas por melhorias, ajustes ou novas configurações no ambiente de TI, enquanto os chamados TOTVS foram relacionados a problemas específicos do ERP utilizado pela empresa.')

def gerar_relatorio_chamados(doc, pipes_data, start_date, end_date):
    """
    Gera a seção 'Relatório de Chamados' com gráfico de linhas.
    """
    if not isinstance(pipes_data, list):
        doc.add_paragraph("Erro: Estrutura dos dados do pipe inválida.")
        return

    # Converter start_date e end_date para date se forem datetime
    start_date = start_date.date() if isinstance(start_date, datetime) else start_date
    end_date = end_date.date() if isinstance(end_date, datetime) else end_date

    total_chamados_abertos = 0
    total_chamados_concluidos = 0
    chamados_por_data = defaultdict(lambda: {'abertos': 0, 'concluidos': 0})

    for card in pipes_data:
        # Acessar os dados diretamente do dicionário
        card_created_date = card['Created at'].date()
        card_finished_date = card['Finished at'].date()
        current_phase = card.get('Current phase')

        if start_date <= card_created_date <= end_date:
            total_chamados_abertos += 1
            chamados_por_data[card_created_date]['abertos'] += 1

        if current_phase == 'Concluído' and start_date <= card_finished_date <= end_date:
            total_chamados_concluidos += 1
            chamados_por_data[card_created_date]['concluidos'] += 1

    doc.add_paragraph('1 - Principais ocorrências')
    doc.add_paragraph(f'Durante o período de {start_date.strftime("%d/%m/%Y")} a {end_date.strftime("%d/%m/%Y")}, o departamento de infraestrutura registrou um total de {total_chamados_abertos} chamados, dos quais {total_chamados_concluidos} foram concluídos com sucesso. Estes chamados incluem solicitações de serviço, incidentes relacionados à operação diária, bem como problemas técnicos. A equipe de infraestrutura conseguiu manter um ritmo constante de atendimento, garantindo a resolução de cada chamado dentro dos prazos estabelecidos, assegurando a disponibilidade e eficiência dos recursos tecnológicos da empresa.')
    doc.add_paragraph("")
    doc.add_paragraph(f'Total de chamados abertos: {total_chamados_abertos}')
    doc.add_paragraph(f'Total de chamados concluídos: {total_chamados_concluidos}')

    grafico_linhas_url = gerar_grafico_linhas(chamados_por_data)
    if grafico_linhas_url:
        img_stream = io.BytesIO(base64.b64decode(grafico_linhas_url))
        doc.add_picture(img_stream, width=Inches(6))

def gerar_graficos_chamados(doc, chamados_por_data):
    """
    Gera gráficos de linha e pizza com base nos dados dos chamados.
    """
    # Gráfico de linha
    grafico_linhas_url = gerar_grafico_linhas(chamados_por_data)
    if grafico_linhas_url:
        img_stream = io.BytesIO(base64.b64decode(grafico_linhas_url))
        doc.add_picture(img_stream, width=Inches(6))

    # Gráfico de pizza
    total_abertos = sum(chamados['abertos'] for chamados in chamados_por_data.values())
    total_concluidos = sum(chamados['concluidos'] for chamados in chamados_por_data.values())
    chamados_resumo = {
        'Chamados Abertos': total_abertos,
        'Chamados Concluídos': total_concluidos
    }
    grafico_pizza_url = gerar_grafico_pizza2(chamados_resumo)
    if grafico_pizza_url:
        img_stream = io.BytesIO(base64.b64decode(grafico_pizza_url))
        doc.add_picture(img_stream, width=Inches(6))

def gerar_chamados_totvs(doc, pipes_data, start_date, end_date):
    """
    Gera a seção 'Chamados da Totvs' com gráfico de colunas.
    """
    if not isinstance(pipes_data, list):
        doc.add_paragraph("Erro: Estrutura dos dados do pipe inválida.")
        return

    start_date = start_date.date() if isinstance(start_date, datetime) else start_date
    end_date = end_date.date() if isinstance(end_date, datetime) else end_date

    total_chamados_abertos = 0
    total_chamados_concluidos = 0

    for card in pipes_data:
        card_created_date = card['Created at'].date()
        current_phase = card.get('Current phase')
        card_title = card.get('Title', 'Descreva sua Solicitação')

        if start_date <= card_created_date <= end_date and 'TOTVS' in card_title.upper():
            total_chamados_abertos += 1
            if current_phase == 'Concluído':
                total_chamados_concluidos += 1

        if start_date <= card_created_date <= end_date and 'DATASUL' in card_title.upper():
            total_chamados_abertos += 1
            if current_phase == 'Concluído':
                total_chamados_concluidos += 1

    doc.add_paragraph('Chamados abertos relacionados à TOTVS')
    doc.add_paragraph(f'Ao longo do período de {start_date.strftime("%d/%m/%Y")} a {end_date.strftime("%d/%m/%Y")}, foram registrados {total_chamados_abertos} chamados relacionados ao sistema TOTVS. A equipe de TI trabalhou com o suporte para resolver os problemas reportados.')
    doc.add_paragraph(f'Total de chamados abertos: {total_chamados_abertos}')
    doc.add_paragraph(f'Total de chamados concluídos: {total_chamados_concluidos}')

    gerar_grafico_colunas_totvs(doc, total_chamados_abertos, total_chamados_concluidos)

def gerar_grafico_colunas_totvs(doc, total_abertos, total_concluidos):
    """
    Gera um gráfico de colunas para comparar chamados abertos e concluídos.
    """
    categorias = ['Chamados Abertos', 'Chamados Concluídos']
    quantidades = [total_abertos, total_concluidos]

    plt.figure(figsize=(6, 4))
    plt.bar(categorias, quantidades, color=['skyblue', 'lightgreen'])
    plt.title('Comparação de Chamados Abertos e Concluídos')
    plt.xlabel('Categorias')
    plt.ylabel('Quantidade')

    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)

    doc.add_picture(img, width=Inches(5))

    plt.close()

def gerar_grafico_linhas(chamados_por_data):
    """
    Gera um gráfico de linhas mostrando a evolução de chamados abertos e concluídos ao longo do tempo.
    """
    datas = sorted(chamados_por_data.keys())
    abertos = [chamados_por_data[data]['abertos'] for data in datas]
    concluidos = [chamados_por_data[data]['concluidos'] for data in datas]

    plt.figure(figsize=(10, 6))
    plt.plot(datas, abertos, label='Chamados Abertos', marker='o')
    plt.plot(datas, concluidos, label='Chamados Concluídos', marker='o')
    plt.xlabel('Data')
    plt.ylabel('Quantidade de Chamados')
    plt.title('Evolução dos Chamados Abertos e Concluídos')
    plt.legend()

    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    grafico_url = base64.b64encode(img.getvalue()).decode()
    plt.close()

    return grafico_url

def gerar_grafico_pizza2(chamados_por_departamento):
    """
    Gera um gráfico de pizza para representar a distribuição de chamados por departamento.
    """
    labels = chamados_por_departamento.keys()
    sizes = chamados_por_departamento.values()

    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')

    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    grafico_url = base64.b64encode(img.getvalue()).decode()
    plt.close()

    return grafico_url

def gerar_grafico_pizza(doc, data):
    """
    Gera gráfico de pizza e insere no documento.
    """
    labels = list(data.keys())
    sizes = list(data.values())

    fig, ax = plt.subplots()
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')

    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(5))
    plt.close()

def gerar_relatorio_com_dados_prometheus(doc, prometheus_data):
    """
    Insere gráficos com os dados retornados do Prometheus.
    """
    for query, result in prometheus_data.items():
        timestamps, values = processar_dados_prometheus(result)
        if timestamps and values:
            doc.add_heading(f'Dados do Prometheus - {query}', level=2)
            gerar_grafico_prometheus(doc, timestamps, values, query)


def processar_dados_prometheus(prometheus_data):
    """
    Processa os dados retornados pelo Prometheus.
    """
    results = prometheus_data.get('data', {}).get('result', [])
    timestamps = []
    values = []

    if results:
        for value in results[0].get('values', []):
            timestamps.append(datetime.fromtimestamp(value[0]))
            values.append(float(value[1]))

    return timestamps, values


def gerar_grafico_prometheus(doc, timestamps, values, title):
    """
    Gera gráfico com dados do Prometheus e insere no relatório.
    """
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=timestamps, y=values, mode='lines', name=title))

    img_stream = BytesIO()
    fig.write_image(img_stream, format='png')
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(5))


def gerar_relatorio_com_dados_grafana(doc, grafana_data):
    """
    Insere gráficos baseados nos dados retornados pelo Grafana.
    """
    panels = grafana_data.get('dashboard', {}).get('panels', [])
    for panel in panels:
        panel_id = panel.get('id')
        panel_title = panel.get('title', 'Título do Painel')
        doc.add_heading(f'Painel: {panel_title}', level=2)
        doc.add_paragraph(f'ID do Painel: {panel_id}')

def adicionar_um_mes(start_date):
    """
    Adiciona um mês ao start_date.
    """
    return start_date + relativedelta(months=1)

def gerar_pagina_inicial(doc, start_date, end_date):
    data_entrega = adicionar_um_mes(start_date)

    p_logo = doc.add_paragraph()
    run_logo = p_logo.add_run()
    run_logo.add_picture('./app/img_itfacil.png', width=Inches(4))
    p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_logo_direita = doc.add_paragraph()
    run_logo_direita = p_logo_direita.add_run()
    run_logo_direita.add_picture('./app/img_ecom.png', width=Inches(4))
    p_logo_direita.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph("Relatório Mensal")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(24)
    run.bold = True

    paragraph = doc.add_paragraph(f'Run BOOK\n{start_date.strftime("%B")}/2024')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(18)

    paragraph = doc.add_paragraph("Relatório Gerencial Mensal\nInformativo – RunBook\nVolumetria e Nível de Serviço")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(12)

    paragraph = doc.add_paragraph("Controle do documento — Versão - 001\nGerente Contrato: Alexandre Bruno")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(10)

    paragraph = doc.add_paragraph(f'Data de entrega – 10{data_entrega.strftime("/%m/%Y")}\nReferência – {start_date.strftime("%d/%m/%Y")} – de 01 até {end_date.strftime("%d/%m/%Y")}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(10)

def gerar_historico_versoes(doc, autorResponsavel, numberVersion):
    """
    Gera a tabela de histórico de versões.
    """
    data_atual = datetime.now().strftime('%d/%m/%Y')

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'DATA'
    hdr_cells[1].text = 'VERSÃO'
    hdr_cells[2].text = 'DESCRIÇÃO DE ALTERAÇÃO'
    hdr_cells[3].text = 'RESPONSÁVEIS'

    row_cells = table.rows[1].cells
    row_cells[0].text = data_atual
    row_cells[1].text = str(numberVersion)
    row_cells[2].text = 'Criação de documento'
    row_cells[3].text = f"Autor: {autorResponsavel}\nRevisor: Leandro Gimenez\nAprovador: Alexandre Bruno"


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Adiciona um hyperlink ao parágrafo.

    :param paragraph: O parágrafo onde será adicionado o hyperlink
    :param url: O URL do link
    :param text: O texto que será mostrado para o link
    :param color: A cor do texto do link (padrão é azul)
    :param underline: Se o link deve ser sublinhado (padrão é True)
    """

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    else:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)

    paragraph._element.append(hyperlink)